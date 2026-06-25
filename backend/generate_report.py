import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()
    
    # --- Title ---
    title = doc.add_heading('VannaQuery Platform', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Project Summary & Feature Analysis Report')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    
    doc.add_paragraph() # Spacing
    
    # --- Executive Summary ---
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        "The VannaQuery Platform is an enterprise-grade, centralized AI-powered database query "
        "intelligence platform. It bridges the gap between natural language and complex database querying, "
        "allowing users and external applications to seamlessly extract data by simply asking questions. "
        "Built on a modern microservices architecture using Docker, it securely connects to various "
        "enterprise databases, translates natural language into SQL using Vanna AI, and safely executes "
        "those queries to return actionable data."
    )
    
    # --- Core Architecture ---
    doc.add_heading('2. Core Architecture', level=1)
    p = doc.add_paragraph("The platform uses a modular, scalable architecture consisting of several interconnected services:")
    
    services = [
        ("Frontend (React)", "A sleek, modern administrative portal hosted on Nginx."),
        ("Backend API (FastAPI)", "The core orchestration layer that handles authentication, database connections, API tokens, and user management."),
        ("Vanna AI Service (FastAPI)", "A dedicated Python microservice that encapsulates the Vanna AI logic and interacts with Large Language Models (LLMs) like OpenAI to translate natural language into SQL."),
        ("PostgreSQL", "The primary application database storing user data, connection profiles, API tokens, and audit logs."),
        ("Redis", "An in-memory data store used for caching and potentially background task brokering."),
        ("Nginx", "A reverse proxy that routes traffic to the frontend and backend services.")
    ]
    
    for name, desc in services:
        bp = doc.add_paragraph(style='List Bullet')
        run = bp.add_run(name + ": ")
        run.bold = True
        bp.add_run(desc)

    # --- Feature Analysis ---
    doc.add_heading('3. Comprehensive Feature Analysis', level=1)
    
    features = [
        ("Natural Language Query Engine", [
            "Text-to-SQL: Users can type questions in plain English, which the system automatically converts into syntactically correct SQL queries tailored to the target database.",
            "Interactive Preview: Before executing the query, the generated SQL is presented to the user for review.",
            "Data Rendering: Query results are returned in a structured format, enabling easy interpretation and integration."
        ]),
        ("Database Connection Management", [
            "Multi-Database Support: The system supports multiple dialects out of the box, including PostgreSQL, MySQL/MariaDB, Microsoft SQL Server, and Oracle Database.",
            "Profile Segregation: Administrators can create distinct connection profiles for different databases or environments, allowing scoped access."
        ]),
        ("API Token Management", [
            "Programmatic Access: Enables external systems and applications to securely query databases via the API.",
            "Granular Controls: Tokens are bound to specific connection profiles, ensuring an application can only query authorized databases.",
            "Rate Limiting: Built-in rate limiting prevents abuse and manages LLM costs."
        ]),
        ("Audit & Observability", [
            "Query Logging: Every generated and executed query is logged with its timestamp, associated user/token, and execution duration.",
            "Performance Metrics: Tracks execution time and row counts, aiding in database performance monitoring.",
            "Accountability: Ensures full traceability of who queried what data and when."
        ]),
        ("Real-Time Dashboard & Settings", [
            "Aggregated Analytics: Provides a high-level overview of system usage, including total queries processed, active API tokens, and system health.",
            "Role-Based Access Control (RBAC): Differentiates between Super Admins, Admins, Users, and API Consumers.",
            "Preferences: Theme customization and profile management."
        ])
    ]
    
    for title, points in features:
        doc.add_heading(f"{title}", level=2)
        for point in points:
            doc.add_paragraph(point, style='List Bullet')

    # --- Security Model ---
    doc.add_heading('4. Security & Compliance Model', level=1)
    doc.add_paragraph("Security is a foundational pillar of the VannaQuery Platform:")
    
    security = [
        ("Authentication", "JWT (JSON Web Tokens) using HS256 algorithm alongside bcrypt password hashing."),
        ("Encryption at Rest", "Sensitive database connection credentials are symmetrically encrypted in the PostgreSQL database using Fernet encryption."),
        ("Query Safety Guardrails", "Built-in AST/Regex filters explicitly block destructive SQL commands (INSERT, UPDATE, DELETE, DROP, ALTER, TRUNCATE) to ensure the system remains strictly read-only."),
        ("Token Security", "API Tokens are generated securely, displayed only once upon creation, and stored as SHA-256 hashes in the database.")
    ]
    
    for name, desc in security:
        bp = doc.add_paragraph(style='List Bullet')
        run = bp.add_run(name + ": ")
        run.bold = True
        bp.add_run(desc)

    # --- Future Roadmap ---
    doc.add_heading('5. Future Roadmap', level=1)
    doc.add_paragraph("The platform is designed with extensibility in mind. Planned future enhancements include:")
    
    roadmap = [
        "Write-back Capabilities: Supporting data modifications via approval workflows.",
        "Row-Level Security (RLS): Enforcing fine-grained data access controls.",
        "Multi-LLM Support: Transitioning from pure OpenAI reliance to supporting local open-source models via Ollama.",
        "BI Integrations: Exporting structured query results directly to BI tools like Power BI and Metabase.",
        "Enterprise SSO: Integration with LDAP and SAML/SSO providers."
    ]
    
    for item in roadmap:
        doc.add_paragraph(item, style='List Bullet')
        
    doc.add_paragraph()
    footer = doc.add_paragraph("Generated for ONGC Vanna Platform")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.italic = True
    footer.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        
    doc.save('/app/VannaQuery_Platform_Report.docx')
    print("Report generated successfully at /app/VannaQuery_Platform_Report.docx")

if __name__ == '__main__':
    create_report()
